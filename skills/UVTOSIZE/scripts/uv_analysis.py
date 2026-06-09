#!/usr/bin/env python3
"""
UV-Vis Quantum Dot Size Analysis Tool
======================================
Reads UV-Vis absorption data from txt files, generates spectra plots,
finds first exciton absorption peaks, and calculates quantum dot sizes
using empirical formulas from the literature.

Supported QD types: PbS, PbSe, CdS, CdSe

References:
  - PbS: Moreels et al., ACS Nano 2009, 3, 3023–3030
  - PbSe: Moreels et al., Chem. Mater. 2007, 19, 6101–6106
  - CdS: Yu et al., Chem. Mater. 2003, 15, 2854–2860
  - CdSe: Yu et al., Chem. Mater. 2003, 15, 2854–2860
"""

import argparse
import os
import sys
import re
import struct
from pathlib import Path
from datetime import datetime

import numpy as np
from scipy.signal import find_peaks, savgol_filter
from scipy.optimize import curve_fit
from scipy.optimize import fsolve
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.ticker import AutoMinorLocator

# ============================================================
# DATA PARSING
# ============================================================

def detect_encoding(filepath):
    """Detect file encoding (UTF-16 LE/BE, UTF-8, GBK, etc.)."""
    with open(filepath, "rb") as f:
        raw = f.read(4)

    # Check for BOM
    if raw[:2] == b"\xff\xfe":
        return "utf-16-le"
    elif raw[:2] == b"\xfe\xff":
        return "utf-16-be"
    elif raw[:3] == b"\xef\xbb\xbf":
        return "utf-8-sig"
    # Try common encodings
    for enc in ["utf-8", "utf-16", "gbk", "gb2312", "latin-1"]:
        try:
            with open(filepath, "r", encoding=enc) as f:
                f.read()
            return enc
        except (UnicodeDecodeError, UnicodeError):
            continue
    return "utf-8"


def parse_uv_txt(filepath):
    """
    Parse a UV-Vis txt data file.

    Expected format:
      Line 1: Sample name (may be quoted)
      Line 2: Column headers (wavelength, absorbance)
      Lines 3+: wavelength,absorbance (comma-separated)

    Returns:
      dict with keys: sample_name, wavelength (np.array), absorbance (np.array)
    """
    encoding = detect_encoding(filepath)
    with open(filepath, "r", encoding=encoding) as f:
        lines = f.readlines()

    # Extract sample name from first line
    first_line = lines[0].strip().strip('"').strip("'")
    sample_name = first_line.replace(" - RawData", "").strip()

    # Parse data lines
    wavelengths = []
    absorbances = []
    for line in lines[1:]:
        line = line.strip().strip('"')
        if not line:
            continue
        parts = line.split(",")
        if len(parts) < 2:
            # Try tab or whitespace
            parts = line.split()
        if len(parts) < 2:
            continue
        try:
            wl = float(parts[0])
            abs_val = float(parts[1])
            wavelengths.append(wl)
            absorbances.append(abs_val)
        except ValueError:
            continue

    if not wavelengths:
        raise ValueError(f"No valid data found in {filepath}")

    wl_array = np.array(wavelengths)
    abs_array = np.array(absorbances)

    # Sort by wavelength
    idx = np.argsort(wl_array)
    wl_array = wl_array[idx]
    abs_array = abs_array[idx]

    return {
        "sample_name": sample_name,
        "filename": os.path.basename(filepath),
        "filepath": filepath,
        "wavelength": wl_array,
        "absorbance": abs_array,
    }


# ============================================================
# PEAK DETECTION
# ============================================================

def find_first_exciton_peak(wavelength, absorbance, qd_type, smooth_window=15, poly_order=3):
    """
    Find the first exciton absorption peak (lowest-energy / longest-wavelength
    significant peak in the spectrum).

    Strategy:
      1. Smooth the spectrum with Savitzky-Golay filter
      2. Find all peaks using scipy.signal.find_peaks
      3. The first exciton peak = the most prominent peak at the longest wavelength
         (lowest energy) that has significant prominence.

    For Pb-chalcogenides (PbS, PbSe), the first exciton peak is typically in the
    NIR region. For Cd-chalcogenides (CdS, CdSe), it's in the visible region.

    Returns:
      dict with: peak_wavelength, peak_absorbance, peak_energy_eV, all_peaks
    """
    # Smooth the spectrum for better peak detection
    n_points = len(absorbance)
    actual_window = min(smooth_window if smooth_window % 2 == 1 else smooth_window + 1,
                        n_points - 2 if (n_points - 2) % 2 == 1 else n_points - 3)
    if actual_window < 5:
        smoothed = absorbance
    else:
        smoothed = savgol_filter(absorbance, actual_window, min(poly_order, actual_window - 2))

    # Find all peaks
    # Set height threshold based on data range
    # Require peaks to have positive absorbance (filter out baseline noise)
    abs_range = np.max(smoothed) - np.min(smoothed)
    height_threshold = max(np.min(smoothed) + 0.05 * abs_range, 0.0)

    peaks, properties = find_peaks(
        smoothed,
        height=height_threshold,
        prominence=0.03 * abs_range,
        distance=5,
        width=3,
    )

    if len(peaks) == 0:
        # Fallback: just find the global maximum
        peak_idx = np.argmax(smoothed)
        peaks = np.array([peak_idx])
        properties = {"peak_heights": [smoothed[peak_idx]],
                       "prominences": [abs_range]}

    # Classify peaks by wavelength region
    # For QDs, the first exciton peak is the significant peak with longest wavelength
    # Filter out noise-level peaks
    significant_peaks = []
    for i, idx in enumerate(peaks):
        wl = wavelength[idx]
        abs_val = absorbance[idx]  # Use original, unsmoothed absorbance
        smoothed_val = smoothed[idx]
        prominence = properties.get("prominences", [0] * len(peaks))[min(i, len(properties.get("prominences", [0])) - 1)]

        significant_peaks.append({
            "index": idx,
            "wavelength": wl,
            "absorbance": abs_val,
            "absorbance_smoothed": smoothed_val,
            "prominence": prominence,
            "energy_eV": 1240.0 / wl if wl > 0 else 0,
        })

    # Sort by wavelength (descending = longest wavelength first)
    significant_peaks.sort(key=lambda p: p["wavelength"], reverse=True)

    # The first exciton peak = the most prominent significant peak
    # at the longest wavelength
    # Strategy: find peaks with substantial prominence, pick the longest-wavelength one
    if significant_peaks:
        # Filter to peaks with at least 10% of max prominence
        max_prominence = max(p["prominence"] for p in significant_peaks)
        good_peaks = [p for p in significant_peaks
                      if p["prominence"] >= 0.1 * max_prominence]

        if good_peaks:
            first_exciton = good_peaks[0]  # Longest wavelength among significant
        else:
            first_exciton = significant_peaks[0]
    else:
        # Absolute fallback
        idx = np.argmax(absorbance)
        first_exciton = {
            "index": idx,
            "wavelength": wavelength[idx],
            "absorbance": absorbance[idx],
            "absorbance_smoothed": absorbance[idx],
            "prominence": abs_range,
            "energy_eV": 1240.0 / wavelength[idx] if wavelength[idx] > 0 else 0,
        }

    return {
        "peak_wavelength": first_exciton["wavelength"],
        "peak_absorbance": first_exciton["absorbance"],
        "peak_energy_eV": first_exciton["energy_eV"],
        "all_peaks": significant_peaks,
        "first_exciton_details": first_exciton,
    }


# ============================================================
# HWHM LITERATURE CONSTRAINTS (per QD type, in meV)
# ============================================================
# Typical HWHM for colloidal QDs with sigma_d < 15%:
#   PbS:  20-120 meV  (Moreels ACS Nano 2009, Weidman ACS Nano 2014)
#   PbSe: 25-150 meV  (Moreels Chem. Mater. 2007, Du ACS Nano 2010)
#   CdS:  40-200 meV  (Yu Chem. Mater. 2003, Vossmeyer JPC 1994)
#   CdSe: 30-150 meV  (Yu Chem. Mater. 2003, Murray Annu. Rev. 2000)
HWHM_CONSTRAINTS = {
    "pbs":  {"hwhm_min_meV": 10,  "hwhm_max_meV": 200},
    "pbse": {"hwhm_min_meV": 15,  "hwhm_max_meV": 250},
    "cds":  {"hwhm_min_meV": 20,  "hwhm_max_meV": 350},
    "cdse": {"hwhm_min_meV": 15,  "hwhm_max_meV": 250},
}

# Default Gaussian fitting window (nm) per QD type
# CdS/CdSe peaks are narrow in wavelength; PbS/PbSe peaks are broad in NIR
DEFAULT_FIT_RANGE = {
    "pbs":  150,   # NIR: ~80 meV at 1500 nm
    "pbse": 200,   # MIR: ~60 meV at 2500 nm
    "cds":   30,   # Vis: ~200 meV at 430 nm
    "cdse":  40,   # Vis: ~170 meV at 550 nm
}


# ============================================================
# GAUSSIAN FITTING & HWHM CALCULATION
# ============================================================

def gaussian(x, A, x0, sigma, baseline):
    """Gaussian function for peak fitting (constant baseline)."""
    return A * np.exp(-(x - x0)**2 / (2 * sigma**2)) + baseline


def gaussian_linear(x, A, x0, sigma, a, b):
    """Gaussian + linear baseline: a*x + b."""
    return A * np.exp(-(x - x0)**2 / (2 * sigma**2)) + a * x + b


def gaussian_exponential(x, A, x0, sigma, c, d, b0):
    """Gaussian + exponential baseline: c*exp(-d*(x-xmin)) + b0.
    Suitable for absorption edges (CdS, CdSe). x_shift = x - x.min() for stability."""
    x_shift = x - x.min()
    return A * np.exp(-(x - x0)**2 / (2 * sigma**2)) + c * np.exp(-d * x_shift) + b0


def gaussian_only(x, A, x0, sigma):
    """Pure Gaussian without baseline (used after baseline subtraction)."""
    return A * np.exp(-(x - x0)**2 / (2 * sigma**2))


def _compute_aic(n, rss, k):
    """Akaike Information Criterion (smaller = better)."""
    if rss <= 0 or n <= k:
        return float("inf")
    return n * np.log(rss / n) + 2 * k


def fit_gaussian_to_peak(wavelength, absorbance, peak_idx, qd_type="pbs",
                          fit_range_nm=None, baseline_mode="auto"):
    """
    Fit a Gaussian to the first exciton peak region and extract HWHM.

    Supports three baseline models:
      - constant:  Gaussian + b
      - linear:    Gaussian + a*x + b
      - exponential: Gaussian + c*exp(-d*(x-xmin)) + b0  (for absorption edges)
      - auto:      Try all three and select best via AIC.
    Literature-based HWHM constraints prevent unphysical fits.

    Parameters:
      wavelength: np.array, wavelengths in nm
      absorbance: np.array, absorbance values
      peak_idx: int, index of the peak in the arrays
      qd_type: str, QD type for HWHM constraint + fit range lookup
      fit_range_nm: float or None, half-width of fitting window in nm.
                    If None, auto-selects based on QD type.

    Returns:
      dict with: hwhm_nm, hwhm_eV, fwhm_nm, fwhm_eV,
                 gaussian_sigma_nm, gaussian_sigma_eV,
                 fit_success, fit_params, fit_r2,
                 fit_wavelength, fit_absorbance
    """
    if fit_range_nm is None:
        fit_range_nm = DEFAULT_FIT_RANGE.get(qd_type, 150)

    peak_wl = wavelength[peak_idx]
    peak_abs = absorbance[peak_idx]

    # Look up HWHM constraints for this QD type
    constraints = HWHM_CONSTRAINTS.get(qd_type, HWHM_CONSTRAINTS["pbs"])
    hwhm_min_meV = constraints["hwhm_min_meV"]
    hwhm_max_meV = constraints["hwhm_max_meV"]

    # Convert energy HWHM limits (meV) to wavelength sigma bounds (nm)
    # Use numerical search for exact conversion since E = 1240/λ is nonlinear
    def _hwhm_meV_to_sigma_nm(target_hwhm_meV):
        """Find sigma_nm that gives target HWHM in meV at peak_wl."""
        target_eV = target_hwhm_meV / 1000.0
        lo, hi = 0.1, fit_range_nm  # search range for sigma
        for _ in range(40):  # binary search
            mid = (lo + hi) / 2
            hwhm_nm_test = mid * np.sqrt(2 * np.log(2))
            if peak_wl - hwhm_nm_test <= 0:
                hi = mid
                continue
            E_high = 1240.0 / (peak_wl - hwhm_nm_test)
            E_low = 1240.0 / (peak_wl + hwhm_nm_test)
            hwhm_eV_test = (E_high - E_low) / 2.0
            if hwhm_eV_test > target_eV:
                hi = mid
            else:
                lo = mid
        return (lo + hi) / 2

    sigma_min_nm = max(1.0, _hwhm_meV_to_sigma_nm(hwhm_min_meV) * 0.5)
    sigma_max_nm = min(fit_range_nm, _hwhm_meV_to_sigma_nm(hwhm_max_meV) * 1.05)

    # Select data within fit_range_nm of the peak
    mask = (wavelength >= peak_wl - fit_range_nm) & (wavelength <= peak_wl + fit_range_nm)
    wl_fit = wavelength[mask]
    abs_fit = absorbance[mask]

    if len(wl_fit) < 10:
        return _fallback_hwhm(peak_wl, peak_abs)

    # Initial guess for Gaussian parameters
    A_guess = max(0.01, peak_abs - np.percentile(abs_fit, 10))
    x0_guess = peak_wl
    sigma_guess = np.clip(fit_range_nm / 4.0, sigma_min_nm, sigma_max_nm)
    b_guess = np.percentile(abs_fit, 10)

    # ---- Define baseline models to try ----
    x_shifted = wl_fit - wl_fit.min()
    baseline_models = {
        "constant": {
            "func": gaussian,
            "p0": [A_guess, x0_guess, sigma_guess, max(b_guess, 0)],
            "bounds": ([0, peak_wl - fit_range_nm, sigma_min_nm, 0],
                        [max(abs_fit) * 2, peak_wl + fit_range_nm, sigma_max_nm,
                         max(abs_fit)]),
            "n_params": 4,
        },
        "linear": {
            "func": gaussian_linear,
            "p0": [A_guess, x0_guess, sigma_guess, 0.0, max(b_guess, 0)],
            "bounds": ([0, peak_wl - fit_range_nm, sigma_min_nm, -0.1, 0],
                        [max(abs_fit) * 2, peak_wl + fit_range_nm, sigma_max_nm,
                         0.1, max(abs_fit)]),
            "n_params": 5,
        },
        "exponential": {
            "func": gaussian_exponential,
            "p0": [A_guess, x0_guess, sigma_guess, A_guess * 0.3, 0.01, max(b_guess * 0.3, 0)],
            "bounds": ([0, peak_wl - fit_range_nm, sigma_min_nm, 0, 0.001, 0],
                        [max(abs_fit) * 2, peak_wl + fit_range_nm, sigma_max_nm,
                         A_guess * 0.5, 0.2, A_guess * 0.5]),
            "n_params": 6,
        },
    }

    def _fit_two_step():
        """Two-step fit: baseline from peak-excluded data, then pure Gaussian.
        Returns (popt_list, r_squared, aic, bl_at_peak, bl_ref) or raises exception on failure.
        popt_list = [A, x0, sigma, c, d, b0] for compatibility with exponential model output.
        bl_ref = reference wavelength for baseline (wl_base.min())."""
        # Step 1: Fit exponential baseline to data OUTSIDE the peak core
        # Exclude peak region: peak_wl +/- 2.5*sigma_guess
        exclude_hw = sigma_guess * 2.5
        baseline_mask = (wl_fit <= peak_wl - exclude_hw) | (wl_fit >= peak_wl + exclude_hw)
        wl_base = wl_fit[baseline_mask]
        abs_base = abs_fit[baseline_mask]

        if len(wl_base) < 6:
            raise ValueError("Not enough data points outside peak for baseline fitting")

        bl_ref = wl_base.min()

        def _exp_baseline(x, c, d, b0):
            return c * np.exp(-d * (x - bl_ref)) + b0

        # Fit baseline
        bl_p0 = [A_guess * 0.3, 0.02, max(b_guess * 0.3, 0)]
        popt_bl, _ = curve_fit(
            _exp_baseline, wl_base, abs_base,
            p0=bl_p0,
            bounds=([0, 0.001, 0], [max(abs_fit), 0.2, max(abs_fit)]),
            maxfev=10000
        )
        c_bl, d_bl, b0_bl = popt_bl

        # Step 2: Subtract baseline from all data in fit window
        baseline_full = c_bl * np.exp(-d_bl * (wl_fit - bl_ref)) + b0_bl
        abs_subtracted = abs_fit - baseline_full

        # Baseline at peak position (for reporting)
        bl_at_peak = c_bl * np.exp(-d_bl * (peak_wl - bl_ref)) + b0_bl

        # Step 3: Fit pure Gaussian to subtracted data
        popt_g, _ = curve_fit(
            gaussian_only, wl_fit, abs_subtracted,
            p0=[max(0.01, peak_abs - bl_at_peak), x0_guess, sigma_guess],
            bounds=([0, peak_wl - fit_range_nm / 2, sigma_min_nm],
                    [max(abs_subtracted) * 2, peak_wl + fit_range_nm / 2, sigma_max_nm]),
            maxfev=10000
        )
        A_g, x0_g, sigma_g = popt_g

        # Combine parameters: [A, x0, sigma, c, d, b0] matching exponential model format
        popt_full = [A_g, x0_g, sigma_g, c_bl, d_bl, b0_bl]

        # Compute fit quality
        fit_full = gaussian_only(wl_fit, A_g, x0_g, sigma_g) + baseline_full
        residuals = abs_fit - fit_full
        ss_res = np.sum(residuals**2)
        ss_tot = np.sum((abs_fit - np.mean(abs_fit))**2)
        r_squared = 1 - ss_res / ss_tot if ss_tot > 0 else 0
        aic = _compute_aic(len(wl_fit), ss_res, 6)  # 6 params: A, x0, sigma, c, d, b0

        return popt_full, r_squared, aic, bl_at_peak, bl_ref

    # Determine which models to try
    if baseline_mode == "auto":
        models_to_try = ["constant", "linear", "two_step", "exponential"]
    elif baseline_mode == "two_step":
        models_to_try = ["two_step"]
    elif baseline_mode in baseline_models:
        models_to_try = [baseline_mode]
    else:
        models_to_try = ["constant"]

    best_result = None
    best_aic = float("inf")

    for model_name in models_to_try:
        if model_name == "two_step":
            # Special handling: two-step baseline subtraction
            try:
                popt, r_squared, aic, bl_at_peak, bl_ref = _fit_two_step()
                peak_frac = popt[0] / (popt[0] + abs(bl_at_peak)) if (popt[0] + abs(bl_at_peak)) > 0 else 0
                penalty = 0 if peak_frac >= 0.3 else (0.3 - peak_frac) * 100
                if bl_at_peak < 0:
                    penalty += abs(bl_at_peak) * 200
                aic_penalized = aic + penalty

                if aic_penalized < best_aic:
                    best_aic = aic_penalized
                    A_fit, x0_fit, sigma_fit = popt[0], popt[1], popt[2]
                    best_result = {
                        "popt": popt,
                        "model_name": "exponential",  # Use exponential for downstream processing
                        "r_squared": r_squared,
                        "aic": aic,
                        "baseline_at_peak": bl_at_peak,
                        "_two_step": True,
                        "_bl_ref": bl_ref,
                    }
            except (RuntimeError, ValueError):
                pass
            continue

        model = baseline_models[model_name]
        try:
            popt, _ = curve_fit(
                model["func"], wl_fit, abs_fit,
                p0=model["p0"], bounds=model["bounds"],
                maxfev=10000
            )
            residuals = abs_fit - model["func"](wl_fit, *popt)
            ss_res = np.sum(residuals**2)
            ss_tot = np.sum((abs_fit - np.mean(abs_fit))**2)
            r_squared = 1 - ss_res / ss_tot if ss_tot > 0 else 0
            aic = _compute_aic(len(wl_fit), ss_res, model["n_params"])

            # Compute baseline at peak for quality check
            if model_name == "constant":
                bl_at_peak = popt[3]
            elif model_name == "linear":
                bl_at_peak = popt[3] * popt[1] + popt[4]
            else:  # exponential
                bl_at_peak = popt[3] * np.exp(-popt[4] * (popt[1] - wl_fit.min())) + popt[5]

            # Penalize models where baseline dominates the peak
            peak_frac = popt[0] / (popt[0] + abs(bl_at_peak)) if (popt[0] + abs(bl_at_peak)) > 0 else 0
            penalty = 0 if peak_frac >= 0.3 else (0.3 - peak_frac) * 100
            # Penalize negative baseline at peak (unphysical: absorbance >= 0)
            if bl_at_peak < 0:
                penalty += abs(bl_at_peak) * 200
            # For linear model, also check baseline across full fit range
            if model_name == "linear":
                bl_full = popt[3] * wl_fit + popt[4]
                bl_min = np.min(bl_full)
                if bl_min < 0:
                    penalty += abs(bl_min) * 200
            aic_penalized = aic + penalty

            if aic_penalized < best_aic:
                best_aic = aic_penalized
                A_fit, x0_fit, sigma_fit = popt[0], popt[1], popt[2]
                best_result = {
                    "popt": popt,
                    "model_name": model_name,
                    "r_squared": r_squared,
                    "aic": aic,
                    "baseline_at_peak": bl_at_peak,
                }
        except (RuntimeError, ValueError):
            continue

    if best_result is None:
        return _fallback_hwhm(peak_wl, peak_abs)

    # Extract best fit results
    popt = best_result["popt"]
    model_name = best_result["model_name"]
    r_squared = best_result["r_squared"]
    is_two_step = best_result.get("_two_step", False)

    A_fit, x0_fit, sigma_fit = popt[0], popt[1], popt[2]

    # Use stored baseline_at_peak from best model selection
    baseline_at_peak = best_result["baseline_at_peak"]
    if model_name == "constant":
        baseline_params = {"b": popt[3]}
    elif model_name == "linear":
        baseline_params = {"a": popt[3], "b": popt[4]}
    else:  # exponential (including two_step which reuses exponential format)
        baseline_params = {"c": popt[3], "d": popt[4], "b0": popt[5]}

    # HWHM in wavelength domain
    hwhm_nm = sigma_fit * np.sqrt(2 * np.log(2))
    E_half_high = 1240.0 / (x0_fit - hwhm_nm)
    E_half_low = 1240.0 / (x0_fit + hwhm_nm)
    hwhm_eV = (E_half_high - E_half_low) / 2.0
    fwhm_nm = 2 * hwhm_nm
    fwhm_eV = 2 * hwhm_eV

    # Generate fitted curves for plotting (extend 1.5×HWHM beyond peak for HWHM marker visibility)
    wl_plot_min = max(wavelength.min(), x0_fit - hwhm_nm * 1.5 - 10)
    wl_plot_max = min(wavelength.max(), x0_fit + hwhm_nm * 1.5 + 10)
    wl_smooth = np.linspace(wl_plot_min, wl_plot_max, 300)

    # Baseline reference point: for two_step it's the baseline data minimum, else wl_fit.min()
    if is_two_step:
        bl_ref = best_result.get("_bl_ref", wl_fit.min())
    else:
        bl_ref = wl_fit.min()

    # Total fit: Gaussian + baseline
    if is_two_step:
        bl_smooth = popt[3] * np.exp(-popt[4] * np.maximum(wl_smooth - bl_ref, 0)) + popt[5]
        abs_smooth = gaussian_only(wl_smooth, A_fit, x0_fit, sigma_fit) + bl_smooth
    else:
        abs_smooth = baseline_models[model_name]["func"](wl_smooth, *popt)

    # Pure Gaussian component (without baseline)
    gauss_only = A_fit * np.exp(-(wl_smooth - x0_fit)**2 / (2 * sigma_fit**2))

    # Baseline component only (clip exponential arg to prevent left-side hook-back)
    if model_name == "constant":
        baseline_only = np.full_like(wl_smooth, baseline_at_peak)
    elif model_name == "linear":
        baseline_only = popt[3] * wl_smooth + popt[4]
    else:  # exponential or two_step
        x_plot_shifted = np.maximum(wl_smooth - bl_ref, 0)
        baseline_only = popt[3] * np.exp(-popt[4] * x_plot_shifted) + popt[5]

    # Post-fit validation
    hwhm_meV_val = hwhm_eV * 1000
    hwhm_warning = None
    if hwhm_meV_val < hwhm_min_meV:
        hwhm_warning = (f"Fitted HWHM ({hwhm_meV_val:.1f} meV) is below the typical "
                        f"range for {qd_type.upper()} QDs ({hwhm_min_meV}-{hwhm_max_meV} meV).")
    elif hwhm_meV_val > hwhm_max_meV:
        hwhm_warning = (f"Fitted HWHM ({hwhm_meV_val:.1f} meV) exceeds the typical "
                        f"range for {qd_type.upper()} QDs ({hwhm_min_meV}-{hwhm_max_meV} meV).")

    half_max = A_fit / 2.0  # Pure Gaussian half-max

    # Compute baseline-subtracted data over fit window
    if model_name == "constant":
        baseline_at_fit = np.full_like(wl_fit, baseline_at_peak)
    elif model_name == "linear":
        baseline_at_fit = popt[3] * wl_fit + popt[4]
    else:  # exponential or two_step
        baseline_at_fit = popt[3] * np.exp(-popt[4] * np.maximum(wl_fit - bl_ref, 0)) + popt[5]
    data_subtracted = abs_fit - baseline_at_fit

    print(f"  Baseline model: {model_name} (AIC={best_aic:.1f})")

    return {
        "hwhm_nm": hwhm_nm,
        "hwhm_eV": hwhm_eV,
        "fwhm_nm": fwhm_nm,
        "fwhm_eV": fwhm_eV,
        "gaussian_sigma_nm": sigma_fit,
        "gaussian_sigma_eV": hwhm_eV / np.sqrt(2 * np.log(2)),
        "fit_success": True,
        "fit_params": {"A": A_fit, "x0": x0_fit, "sigma": sigma_fit,
                        "baseline": baseline_at_peak,
                        "baseline_model": model_name,
                        "baseline_params": baseline_params,
                        "half_max": half_max},
        "fit_r2": r_squared,
        "fit_wavelength": wl_smooth,
        "fit_absorbance": abs_smooth,
        "gauss_only": gauss_only,
        "baseline_only": baseline_only,
        "baseline_model": model_name,
        "hwhm_warning": hwhm_warning,
        "hwhm_min_meV": hwhm_min_meV,
        "hwhm_max_meV": hwhm_max_meV,
        "aic": best_aic,
        "fit_wl_data": wl_fit,
        "abs_subtracted": data_subtracted,
    }


def _fallback_hwhm(peak_wl, peak_abs):
    """Fallback HWHM estimate when Gaussian fitting fails."""
    return {
        "hwhm_nm": np.nan,
        "hwhm_eV": np.nan,
        "fwhm_nm": np.nan,
        "fwhm_eV": np.nan,
        "gaussian_sigma_nm": np.nan,
        "gaussian_sigma_eV": np.nan,
        "fit_success": False,
        "fit_params": None,
        "fit_r2": None,
        "fit_wavelength": None,
        "fit_absorbance": None,
        "gauss_only": None,
        "baseline_only": None,
        "baseline_model": None,
        "fit_wl_data": None,
        "abs_subtracted": None,
        "hwhm_warning": None,
        "hwhm_min_meV": None,
        "hwhm_max_meV": None,
    }


def calculate_size_distribution(qd_type, diameter_nm, hwhm_eV, peak_wavelength_nm):
    """
    Calculate the size distribution width sigma (σ) from the HWHM of the
    first exciton peak using the error propagation method:

        σ_d = HWHM(E) / |dE/dd|_{d=d_mean}

    where dE/dd is the derivative of the sizing curve at the mean diameter.

    Reference:
      Nikolaev & Averkiev, Appl. Phys. Lett. 95, 263107 (2009)
      Wu et al., Appl. Phys. Lett. 51, 710 (1987)

    Returns:
      dict with sigma_nm, relative_sigma_percent, deriv_dE_dd, hwhm_eV, diameter_nm
    """
    qd_type = qd_type.lower()

    # Calculate |dE/dd| at the mean diameter
    dE_dd = abs(calculate_dE_dd(qd_type, diameter_nm, peak_wavelength_nm))

    if dE_dd < 1e-12:
        return {
            "sigma_nm": np.nan,
            "relative_sigma_percent": np.nan,
            "deriv_dE_dd": dE_dd,
            "hwhm_eV": hwhm_eV,
            "diameter_nm": diameter_nm,
        }

    sigma_nm = hwhm_eV / dE_dd
    relative_sigma = (sigma_nm / diameter_nm) * 100.0 if diameter_nm > 0 else np.nan

    return {
        "sigma_nm": sigma_nm,
        "relative_sigma_percent": relative_sigma,
        "deriv_dE_dd": dE_dd,
        "hwhm_eV": hwhm_eV,
        "diameter_nm": diameter_nm,
    }


def calculate_dE_dd(qd_type, d, peak_wavelength_nm):
    """
    Calculate dE/dd — the derivative of exciton energy with respect to
    quantum dot diameter — evaluated at diameter d.

    For PbS and PbSe (energy-based sizing curves):
      E(d) = E_bulk + 1/(a*d² + b*d + c)
      dE/dd = -(2a*d + b) / (a*d² + b*d + c)²

    For CdS and CdSe (wavelength-based sizing curves):
      D(λ) = polynomial
      dD/dλ = polynomial derivative
      E = 1240/λ
      dE/dD = dE/dλ × dλ/dD = (-1240/λ²) / (dD/dλ)
    """
    if qd_type == "pbs":
        a, b = 0.0252, 0.283
        u = a * d**2 + b * d
        return -(2 * a * d + b) / (u**2)

    elif qd_type == "pbse":
        a, b, c = 0.016, 0.209, 0.45
        u = a * d**2 + b * d + c
        return -(2 * a * d + b) / (u**2)

    elif qd_type == "cds":
        lam = peak_wavelength_nm
        # dD/dλ = -1.99563e-7*λ² + 3.9114e-4*λ - 9.2352e-2
        dD_dlam = -1.99563e-7 * lam**2 + 3.9114e-4 * lam - 9.2352e-2
        if abs(dD_dlam) < 1e-12:
            return 0.0
        dE_dD = (-1240.0 / lam**2) / dD_dlam
        return dE_dD

    elif qd_type == "cdse":
        lam = peak_wavelength_nm
        # dD/dλ = 6.4488e-9*λ³ - 7.9725e-6*λ² + 3.2484e-3*λ - 0.4277
        dD_dlam = 6.4488e-9 * lam**3 - 7.9725e-6 * lam**2 + 3.2484e-3 * lam - 0.4277
        if abs(dD_dlam) < 1e-12:
            return 0.0
        dE_dD = (-1240.0 / lam**2) / dD_dlam
        return dE_dD

    else:
        return 0.0


# ============================================================
# SIZE CALCULATION FORMULAS
# ============================================================

def calculate_size_pbs(peak_wavelength_nm):
    """
    PbS sizing formula from Moreels et al., ACS Nano 2009, 3, 3023-3030.

    E0 = 0.41 + 1 / (0.0252 * d^2 + 0.283 * d)

    where:
      E0 = first exciton energy in eV
      d  = quantum dot diameter in nm
      0.41 eV = bulk PbS bandgap

    Solving for d:
      0.0252*d^2 + 0.283*d - 1/(E0 - 0.41) = 0
      d = (-0.283 + sqrt(0.283^2 + 4*0.0252/(E0-0.41))) / (2*0.0252)
    """
    E0 = 1240.0 / peak_wavelength_nm

    if E0 <= 0.41:
        raise ValueError(
            f"Peak energy {E0:.3f} eV is below the PbS bulk bandgap (0.41 eV). "
            f"The peak at {peak_wavelength_nm:.1f} nm may not be the first exciton peak, "
            f"or the quantum dot is extremely large (> bulk-like)."
        )

    a = 0.0252
    b = 0.283
    c = -1.0 / (E0 - 0.41)

    discriminant = b**2 - 4 * a * c
    if discriminant < 0:
        raise ValueError(
            f"Cannot solve sizing equation: negative discriminant. "
            f"Check that the peak wavelength {peak_wavelength_nm:.1f} nm is correct."
        )

    d = (-b + np.sqrt(discriminant)) / (2 * a)
    return d, E0


def calculate_size_pbse(peak_wavelength_nm):
    """
    PbSe sizing formula from Moreels et al., Chem. Mater. 2007, 19, 6101-6106.

    E0 = 0.278 + 1 / (0.016 * d^2 + 0.209 * d + 0.45)

    where:
      E0 = first exciton energy in eV
      d  = quantum dot diameter in nm
      0.278 eV = bulk PbSe bandgap

    Solving for d:
      0.016*d^2 + 0.209*d + 0.45 - 1/(E0-0.278) = 0
      d = (-0.209 + sqrt(0.209^2 - 4*0.016*(0.45 - 1/(E0-0.278)))) / (2*0.016)
    """
    E0 = 1240.0 / peak_wavelength_nm

    if E0 <= 0.278:
        raise ValueError(
            f"Peak energy {E0:.3f} eV is below the PbSe bulk bandgap (0.278 eV). "
            f"The peak at {peak_wavelength_nm:.1f} nm may not be the first exciton peak."
        )

    a = 0.016
    b = 0.209
    c = 0.45 - 1.0 / (E0 - 0.278)

    discriminant = b**2 - 4 * a * c
    if discriminant < 0:
        raise ValueError(
            f"Cannot solve sizing equation: negative discriminant. "
            f"Check that the peak wavelength {peak_wavelength_nm:.1f} nm is correct."
        )

    d = (-b + np.sqrt(discriminant)) / (2 * a)
    return d, E0


def calculate_size_cds(peak_wavelength_nm):
    """
    CdS sizing formula from Yu et al., Chem. Mater. 2003, 15, 2854-2860.

    D = -6.6521e-8 * lambda^3 + 1.9557e-4 * lambda^2
        - 9.2352e-2 * lambda + 13.29

    where:
      D = quantum dot diameter in nm
      lambda = first exciton peak wavelength in nm

    Valid range: ~360-500 nm (1-6 nm QDs)
    """
    l = peak_wavelength_nm
    D = (-6.6521e-8 * l**3
         + 1.9557e-4 * l**2
         - 9.2352e-2 * l
         + 13.29)
    E0 = 1240.0 / peak_wavelength_nm
    return D, E0


def calculate_size_cdse(peak_wavelength_nm):
    """
    CdSe sizing formula from Yu et al., Chem. Mater. 2003, 15, 2854-2860.

    D = 1.6122e-9 * lambda^4 - 2.6575e-6 * lambda^3
        + 1.6242e-3 * lambda^2 - 0.4277 * lambda + 41.57

    where:
      D = quantum dot diameter in nm
      lambda = first exciton peak wavelength in nm

    Valid range: ~510-640 nm (1.5-8 nm QDs)
    """
    l = peak_wavelength_nm
    D = (1.6122e-9 * l**4
         - 2.6575e-6 * l**3
         + 1.6242e-3 * l**2
         - 0.4277 * l
         + 41.57)
    E0 = 1240.0 / peak_wavelength_nm
    return D, E0


# Formula registry
SIZE_FORMULAS = {
    "pbs": {
        "func": calculate_size_pbs,
        "name": "PbS",
        "reference": "Moreels et al., ACS Nano 2009, 3, 3023–3030",
        "formula_str": "E₀ = 0.41 + 1 / (0.0252·d² + 0.283·d)",
        "deriv_str": "dE/dd = -(0.0504d + 0.283) / (0.0252d² + 0.283d)²",
        "variable": "E₀ (eV) → d (nm)",
        "method": "Solve quadratic equation for d",
        "valid_range": "~730–2500 nm (0.5–1.7 eV), 2.3–10 nm",
        "bulk_bandgap_eV": 0.41,
    },
    "pbse": {
        "func": calculate_size_pbse,
        "name": "PbSe",
        "reference": "Moreels et al., Chem. Mater. 2007, 19, 6101–6106",
        "formula_str": "E₀ = 0.278 + 1 / (0.016·d² + 0.209·d + 0.45)",
        "deriv_str": "dE/dd = -(0.032d + 0.209) / (0.016d² + 0.209d + 0.45)²",
        "variable": "E₀ (eV) → d (nm)",
        "method": "Solve quadratic equation for d",
        "valid_range": "~800–4000 nm (0.3–1.5 eV), 2–20 nm",
        "bulk_bandgap_eV": 0.278,
    },
    "cds": {
        "func": calculate_size_cds,
        "name": "CdS",
        "reference": "Yu et al., Chem. Mater. 2003, 15, 2854–2860",
        "formula_str": "D = -6.6521×10⁻⁸·λ³ + 1.9557×10⁻⁴·λ² - 9.2352×10⁻²·λ + 13.29",
        "deriv_str": "dD/dλ = -1.9956×10⁻⁷·λ² + 3.9114×10⁻⁴·λ - 9.2352×10⁻²; dE/dD = (-1240/λ²) / (dD/dλ)",
        "variable": "λ (nm) → D (nm)",
        "method": "Direct polynomial evaluation",
        "valid_range": "~360–500 nm, 1–6 nm",
        "bulk_bandgap_eV": 2.42,
    },
    "cdse": {
        "func": calculate_size_cdse,
        "name": "CdSe",
        "reference": "Yu et al., Chem. Mater. 2003, 15, 2854–2860",
        "formula_str": "D = 1.6122×10⁻⁹·λ⁴ - 2.6575×10⁻⁶·λ³ + 1.6242×10⁻³·λ² - 0.4277·λ + 41.57",
        "deriv_str": "dD/dλ = 6.4488×10⁻⁹·λ³ - 7.9725×10⁻⁶·λ² + 3.2484×10⁻³·λ - 0.4277; dE/dD = (-1240/λ²) / (dD/dλ)",
        "variable": "λ (nm) → D (nm)",
        "method": "Direct polynomial evaluation",
        "valid_range": "~510–640 nm, 1.5–8 nm",
        "bulk_bandgap_eV": 1.74,
    },
}


VALID_RANGES = {
    "pbs": {"wl_min": 730, "wl_max": 2500, "size_min": 2.3, "size_max": 10.0,
            "eV_min": 0.50, "eV_max": 1.70},
    "pbse": {"wl_min": 800, "wl_max": 4000, "size_min": 2.0, "size_max": 20.0,
             "eV_min": 0.31, "eV_max": 1.55},
    "cds": {"wl_min": 360, "wl_max": 500, "size_min": 1.0, "size_max": 6.0,
            "eV_min": 2.48, "eV_max": 3.44},
    "cdse": {"wl_min": 510, "wl_max": 640, "size_min": 1.5, "size_max": 8.0,
             "eV_min": 1.94, "eV_max": 2.43},
}


def check_validity(qd_type, peak_wl_nm, size_nm):
    """Check if results fall within the calibrated range. Returns list of warnings."""
    qd_type = qd_type.lower()
    warnings = []
    r = VALID_RANGES.get(qd_type)
    if r is None:
        return warnings

    eV = 1240.0 / peak_wl_nm
    if peak_wl_nm < r["wl_min"] or peak_wl_nm > r["wl_max"]:
        warnings.append(
            f"[WARNING] Peak wavelength {peak_wl_nm:.1f} nm ({eV:.3f} eV) is outside the "
            f"calibrated range of {r['wl_min']:.0f}-{r['wl_max']:.0f} nm "
            f"({r['eV_min']:.2f}-{r['eV_max']:.2f} eV) for {qd_type.upper()}. "
            f"The detected peak may not be the first exciton transition."
        )
    if size_nm < r["size_min"] or size_nm > r["size_max"]:
        warnings.append(
            f"[WARNING] Calculated size {size_nm:.3f} nm is outside the calibrated range "
            f"of {r['size_min']:.1f}-{r['size_max']:.1f} nm for {qd_type.upper()}. "
            f"Verify the peak assignment or use TEM/XRD for confirmation."
        )
    return warnings


def calculate_size(qd_type, peak_wavelength_nm):
    """Calculate QD size using the appropriate formula.
    Returns (size_nm, energy_eV, warnings_list)."""
    qd_type = qd_type.lower()
    if qd_type not in SIZE_FORMULAS:
        raise ValueError(f"Unknown QD type: {qd_type}. Supported: {list(SIZE_FORMULAS.keys())}")
    size_nm, E0 = SIZE_FORMULAS[qd_type]["func"](peak_wavelength_nm)
    warnings = check_validity(qd_type, peak_wavelength_nm, size_nm)
    return size_nm, E0, warnings


# ============================================================
# PLOTTING
# ============================================================

# ============================================================
# Colorblind-friendly palette (Wong 2011, Nature Methods)
# ============================================================
CB_PALETTE = {
    "blue":       "#0072B2",
    "orange":     "#E69F00",
    "green":      "#009E73",
    "vermilion":  "#D55E00",
    "purple":     "#CC79A7",
    "yellow":     "#F0E442",
    "sky_blue":   "#56B4E9",
    "black":      "#000000",
}

# Figure width presets (cm → inches)
FIGURE_WIDTHS = {
    "single": 8.5 / 2.54,    # ~3.35"
    "double": 17.0 / 2.54,   # ~6.69"
    "full":   12.0,           # legacy default
}


def setup_publication_style():
    """Configure matplotlib rcParams for publication-quality output."""
    # Fonts: Times New Roman (serif) with Arial/SimSun fallback
    matplotlib.rcParams.update({
        "font.family":          "serif",
        "font.serif":           ["Times New Roman", "STIXGeneral", "DejaVu Serif"],
        "font.sans-serif":      ["Arial", "SimSun", "Microsoft YaHei", "DejaVu Sans"],
        "mathtext.fontset":     "stix",
        "mathtext.default":     "regular",
        "font.size":            10,
        "axes.titlesize":       12,
        "axes.labelsize":       11,
        "xtick.labelsize":      9,
        "ytick.labelsize":      9,
        "legend.fontsize":      9,
        "figure.dpi":           300,
        "savefig.dpi":          300,
        "savefig.bbox":         "tight",
        "savefig.pad_inches":   0.05,
        # Lines
        "lines.linewidth":      1.2,
        "lines.markersize":     6,
        # Ticks — inward
        "xtick.direction":      "in",
        "ytick.direction":      "in",
        "xtick.major.size":     4.5,
        "ytick.major.size":     4.5,
        "xtick.minor.size":     2.5,
        "ytick.minor.size":     2.5,
        "xtick.major.width":    0.8,
        "ytick.major.width":    0.8,
        "xtick.minor.width":    0.6,
        "ytick.minor.width":    0.6,
        "xtick.top":            True,
        "ytick.right":          True,
        # Axes
        "axes.linewidth":       0.8,
        "axes.grid":            True,
        "grid.alpha":           0.25,
        "grid.linestyle":       "--",
        "grid.linewidth":       0.4,
        # Legend
        "legend.frameon":       True,
        "legend.framealpha":    0.85,
        "legend.edgecolor":     "#cccccc",
        "legend.fancybox":      False,
        "legend.borderpad":     0.3,
        "legend.borderaxespad": 0.5,
    })


def create_uv_plot(data, peak_info, qd_type, size_nm, output_dir,
                    gaussian_fit=None, sigma_info=None,
                    x_range=None, y_range=None,
                    output_formats=None, dpi=300,
                    figure_width="double",
                    show_title=True, show_legend=True,
                    show_annotation=False):
    """
    Generate a publication-quality UV-Vis spectrum plot.

    Parameters
    ----------
    data : dict
        Parsed UV data (wavelength, absorbance, sample_name).
    peak_info : dict
        First exciton peak information.
    qd_type : str
        Quantum dot type (pbs/pbse/cds/cdse).
    size_nm : float
        Calculated quantum dot diameter in nm.
    output_dir : str
        Directory to save output files.
    gaussian_fit : dict or None
        Gaussian fitting results.
    sigma_info : dict or None
        Size distribution results.
    x_range : tuple or None
        Custom X-axis range (xmin, xmax) in nm. Auto if None.
    y_range : tuple or None
        Custom Y-axis range (ymin, ymax). Auto if None.
    output_formats : list or None
        Output formats: ['png', 'pdf', 'svg']. Default: ['png'].
    dpi : int
        Resolution for PNG output (300 or 600).
    figure_width : str or float
        Figure width: "single" (8.5 cm), "double" (17 cm), or custom float in inches.
    show_title : bool
        Whether to show the plot title.
    show_legend : bool
        Whether to show the legend.

    Returns
    -------
    dict
        Mapping of format → file path, e.g. {'png': '...', 'pdf': '...', 'svg': '...'}.
    """
    setup_publication_style()

    wl = data["wavelength"]
    abs_data = data["absorbance"]
    peak_wl = peak_info["peak_wavelength"]
    peak_abs = peak_info["peak_absorbance"]

    formula_info = SIZE_FORMULAS.get(qd_type.lower(), {})
    qd_name = formula_info.get("name", qd_type.upper())

    # Resolve figure width
    if isinstance(figure_width, str):
        fig_w = FIGURE_WIDTHS.get(figure_width, FIGURE_WIDTHS["double"])
    else:
        fig_w = float(figure_width)
    fig_h = fig_w * 0.618  # Golden ratio aspect

    if output_formats is None:
        output_formats = ["png"]

    fig, ax = plt.subplots(figsize=(fig_w, fig_h))

    # ---- Traces ----
    # 1. Absorption spectrum (blue)
    ax.plot(wl, abs_data, linewidth=0.8, color=CB_PALETTE["blue"],
            label="Absorption", zorder=2)

    # 2. Fit components
    if gaussian_fit and gaussian_fit["fit_success"]:
        # 2. Pure Gaussian fit (solid, vermilion)
        ax.plot(gaussian_fit["fit_wavelength"], gaussian_fit["gauss_only"],
                linewidth=1.0, color=CB_PALETTE["vermilion"], linestyle="-",
                label="Pure Gauss Fit", zorder=3.5)

        # 3. HWHM calibration on pure Gaussian (green, half-max = A/2)
        wl0 = gaussian_fit["fit_params"]["x0"]
        hwhm_nm = gaussian_fit["hwhm_nm"]
        hwhm_eV = gaussian_fit["hwhm_eV"]
        # HWHM on pure Gaussian: half-max = A/2
        half_max = gaussian_fit["fit_params"]["A"] / 2.0
        wl_left = wl0 - hwhm_nm
        wl_right = wl0 + hwhm_nm

        # HWHM dashed line at half-max
        ax.hlines(y=half_max, xmin=wl_left, xmax=wl_right,
                   colors=CB_PALETTE["green"], linewidth=1.0, linestyle="--",
                   label="HWHM",
                   zorder=4)
        # Vertical end ticks
        tick_h = (abs_data.max() - abs_data.min()) * 0.025
        ax.vlines(x=[wl_left, wl_right],
                   ymin=half_max - tick_h, ymax=half_max + tick_h,
                   colors=CB_PALETTE["green"], linewidth=1.0, zorder=5)

    # 4. Peak marker (vermilion)
    ax.scatter([peak_wl], [peak_abs], color=CB_PALETTE["vermilion"], s=50, zorder=6,
               edgecolors="white", linewidths=0.8)
    ax.axvline(x=peak_wl, color=CB_PALETTE["vermilion"], linestyle="--",
               alpha=0.35, linewidth=0.8, zorder=1)

    # ---- Annotation box (optional) ----
    if show_annotation:
        ann_lines = [
            f"λ = {peak_wl:.1f} nm",
            f"E = {1240/peak_wl:.3f} eV",
            f"d = {size_nm:.2f} nm",
        ]
        if gaussian_fit and gaussian_fit["fit_success"]:
            ann_lines.append(f"HWHM = {gaussian_fit['hwhm_nm']:.1f} nm")
            ann_lines.append(f"R² = {gaussian_fit['fit_r2']:.4f}")
        if sigma_info and not np.isnan(sigma_info.get("sigma_nm", np.nan)):
            ann_lines.append(f"σ = {sigma_info['sigma_nm']:.3f} nm ({sigma_info['relative_sigma_percent']:.1f}%)")

        annotation_text = "\n".join(ann_lines)

        wl_range_ann = wl.max() - wl.min()
        mid_wl = (wl.min() + wl.max()) / 2
        peak_is_right = peak_wl > mid_wl

        if peak_is_right:
            text_x = wl.min() + 0.05 * wl_range_ann
            ha = "left"
        else:
            text_x = wl.max() - 0.05 * wl_range_ann
            ha = "right"

        text_y = abs_data.max() * 0.88

        ax.annotate(
            annotation_text,
            xy=(peak_wl, peak_abs),
            xytext=(text_x, text_y),
            horizontalalignment=ha,
            bbox=dict(boxstyle="round,pad=0.5", facecolor="white",
                       edgecolor="#cccccc", alpha=0.90, linewidth=0.5),
            fontsize=7.5,
            fontfamily="sans-serif",
            zorder=99,
            arrowprops=dict(arrowstyle="->", color="#999999", lw=0.8,
                            connectionstyle="arc3,rad=0.2"),
        )

    # ---- Labels ----
    ax.set_xlabel("Wavelength (nm)", fontweight="normal")
    ax.set_ylabel("Absorbance (a.u.)", fontweight="normal")
    if show_title:
        ax.set_title(
            f"UV-Vis Absorption — {qd_name} QDs ({data['sample_name']})",
            fontweight="bold", pad=8
        )

    # ---- Legend ----
    if show_legend:
        legend = ax.legend(
            loc="best",
            frameon=True,
            fancybox=False,
            edgecolor="#cccccc",
            borderpad=0.3,
            borderaxespad=0.5,
        )
        # Make legend non-occluding: place outside if it overlaps data
        # Try "best" first; matplotlib auto-optimizes to avoid data

    # ---- Ticks ----
    ax.xaxis.set_minor_locator(AutoMinorLocator(4))
    ax.yaxis.set_minor_locator(AutoMinorLocator(4))
    ax.tick_params(axis="both", which="major", direction="in",
                   top=True, right=True, labelsize=9, width=0.8, length=4.5)
    ax.tick_params(axis="both", which="minor", direction="in",
                   top=True, right=True, width=0.6, length=2.5)

    # ---- Axis limits ----
    wl_range = wl.max() - wl.min()

    # X-axis: use custom if provided, otherwise auto with 3% padding
    pad_x = 0.03 * wl_range
    x_min_auto = wl.min() - pad_x
    x_max_auto = wl.max() + pad_x
    if x_range is not None and len(x_range) == 2:
        ax.set_xlim(
            float(x_range[0]) if x_range[0] is not None else x_min_auto,
            float(x_range[1]) if x_range[1] is not None else x_max_auto,
        )
    else:
        ax.set_xlim(x_min_auto, x_max_auto)

    # Y-axis: use custom if provided, otherwise auto with 3%/10% padding
    abs_range = abs_data.max() - abs_data.min()
    y_min_auto = min(0, abs_data.min() - 0.03 * abs_range)
    y_max_auto = abs_data.max() + 0.10 * abs_range
    if y_range is not None and len(y_range) == 2:
        ax.set_ylim(
            float(y_range[0]) if y_range[0] is not None else y_min_auto,
            float(y_range[1]) if y_range[1] is not None else y_max_auto,
        )
    else:
        ax.set_ylim(y_min_auto, y_max_auto)

    plt.tight_layout(pad=0.8)

    # ---- Save in requested formats ----
    safe_name = re.sub(r'[^\w\s-]', '', data["sample_name"]).strip()
    safe_name = re.sub(r'[-\s]+', '_', safe_name)
    base = os.path.join(output_dir, f"UV_spectrum_{safe_name}")

    paths = {}

    for fmt in output_formats:
        fmt_lower = fmt.lower().strip()
        if fmt_lower == "png":
            fpath = base + ".png"
            plt.savefig(fpath, dpi=dpi, bbox_inches="tight",
                        facecolor="white", edgecolor="none")
            paths["png"] = fpath
        elif fmt_lower == "pdf":
            fpath = base + ".pdf"
            plt.savefig(fpath, format="pdf", bbox_inches="tight",
                        facecolor="white", edgecolor="none")
            paths["pdf"] = fpath
        elif fmt_lower == "svg":
            fpath = base + ".svg"
            plt.savefig(fpath, format="svg", bbox_inches="tight",
                        facecolor="white", edgecolor="none")
            paths["svg"] = fpath

    plt.close(fig)

    # Return dict with all output format paths
    return paths


# ============================================================
# WORD DOCUMENT GENERATION
# ============================================================

def create_word_report(data, peak_info, qd_type, size_nm, plot_path, output_dir,
                       warnings=None, gaussian_fit=None, sigma_info=None):
    """Generate a detailed Word document with the analysis results and formulas."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        print("python-docx not installed. Skipping Word report generation.")
        return None

    formula_info = SIZE_FORMULAS.get(qd_type.lower(), {})
    qd_name = formula_info.get("name", qd_type.upper())

    doc = Document()

    # --- Styles ---
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # --- Title ---
    title = doc.add_heading(f"UV-Vis Analysis Report — {qd_name} Quantum Dots", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
                      style="Normal").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # --- Section 1: Sample Information ---
    doc.add_heading("1. Sample Information", level=1)
    info_table = doc.add_table(rows=5, cols=2, style="Light Grid Accent 1")
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("Sample Name", data["sample_name"]),
        ("Source File", data["filename"]),
        ("Quantum Dot Type", qd_name),
        ("Measurement Type", "UV-Vis Absorption Spectroscopy"),
        ("Wavelength Range", f"{data['wavelength'].min():.0f} – {data['wavelength'].max():.0f} nm"),
    ]
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = str(value)
        for cell in info_table.rows[i].cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

    # --- Section 2: Peak Analysis ---
    doc.add_heading("2. First Exciton Peak Analysis", level=1)

    doc.add_paragraph(
        "The first exciton absorption peak is identified as the most prominent "
        "absorption feature at the longest wavelength (lowest energy) in the UV-Vis "
        "spectrum. This peak corresponds to the 1S(e)–1S(h) electronic transition "
        "and its position directly relates to the quantum dot size via quantum "
        "confinement effects."
    )

    peak_table = doc.add_table(rows=5, cols=2, style="Light Grid Accent 1")
    peak_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    peak_wl = peak_info["peak_wavelength"]
    peak_eV = 1240.0 / peak_wl
    peak_data = [
        ("Peak Wavelength (λ)", f"{peak_wl:.2f} nm"),
        ("Peak Absorbance", f"{peak_info['peak_absorbance']:.4f} a.u."),
        ("Peak Energy (E₀)", f"{peak_eV:.3f} eV"),
        ("Calculation Method", "Savitzky-Golay smoothing + scipy.signal.find_peaks"),
        ("Number of Peaks Detected", str(len(peak_info.get("all_peaks", [])))),
    ]
    for i, (label, value) in enumerate(peak_data):
        peak_table.rows[i].cells[0].text = label
        peak_table.rows[i].cells[1].text = str(value)
        for cell in peak_table.rows[i].cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)

    # List all detected peaks
    all_peaks = peak_info.get("all_peaks", [])
    if len(all_peaks) > 1:
        doc.add_paragraph()
        doc.add_paragraph("All detected absorption peaks:", style="Normal").bold = True
        for i, p in enumerate(all_peaks):
            marker = " ★ First Exciton" if abs(p["wavelength"] - peak_wl) < 0.01 else ""
            doc.add_paragraph(
                f"  Peak {i+1}: λ = {p['wavelength']:.1f} nm, "
                f"A = {p['absorbance']:.4f}, "
                f"E = {1240/p['wavelength']:.3f} eV"
                f"{marker}",
                style="Normal"
            )

    doc.add_paragraph()

    # --- Section 3: Size Calculation ---
    doc.add_heading("3. Quantum Dot Size Calculation", level=1)

    doc.add_heading("3.1 Empirical Formula", level=2)

    doc.add_paragraph(
        f"The following empirical sizing formula was used for {qd_name} quantum dots:"
    )

    # Formula display
    formula_para = doc.add_paragraph()
    formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_run = formula_para.add_run(formula_info["formula_str"])
    formula_run.font.size = Pt(12)
    formula_run.font.italic = True
    formula_run.font.name = "Times New Roman"

    doc.add_paragraph()
    doc.add_paragraph(f"Reference: {formula_info['reference']}")
    doc.add_paragraph(f"Variable convention: {formula_info['variable']}")
    doc.add_paragraph(f"Solution method: {formula_info['method']}")
    doc.add_paragraph(f"Valid range: {formula_info['valid_range']}")
    doc.add_paragraph(f"Bulk bandgap (E_g): {formula_info['bulk_bandgap_eV']} eV")

    doc.add_paragraph()

    doc.add_heading("3.2 Calculation Steps", level=2)

    if qd_type.lower() in ("pbs", "pbse"):
        # Show detailed steps for Pb-chalcogenides
        E0 = 1240.0 / peak_wl
        bg = formula_info["bulk_bandgap_eV"]

        steps = [
            f"Step 1: Convert peak wavelength to energy",
            f"  E₀ = 1240 / λ = 1240 / {peak_wl:.2f} = {E0:.4f} eV",
            "",
            f"Step 2: Apply sizing formula",
            f"  {formula_info['formula_str']}",
            f"  E₀ – E_g(bulk) = {E0:.4f} – {bg:.4f} = {E0 - bg:.4f} eV",
            f"  1 / (E₀ – E_g) = 1 / {E0 - bg:.4f} = {1/(E0-bg):.4f}",
            "",
            f"Step 3: Solve quadratic equation for d",
        ]

        if qd_type.lower() == "pbs":
            a, b_val = 0.0252, 0.283
            c_val = -1.0 / (E0 - 0.41)
            steps.append(f"  0.0252·d² + 0.283·d - {1/(E0-0.41):.4f} = 0")
        else:
            a, b_val = 0.016, 0.209
            c_val = 0.45 - 1.0 / (E0 - 0.278)
            steps.append(f"  0.016·d² + 0.209·d + ({c_val:.4f}) = 0")

        disc = b_val**2 - 4*a*c_val
        d_pos = (-b_val + np.sqrt(disc)) / (2*a)

        steps += [
            f"  Discriminant Δ = {b_val}² – 4×{a}×({c_val:.4f}) = {disc:.6f}",
            f"  d = (–{b_val} + √{disc:.6f}) / (2×{a})",
            f"  d = {d_pos:.4f} nm",
        ]

        for step in steps:
            p = doc.add_paragraph(step, style="Normal")
            if step.startswith("Step"):
                p.runs[0].bold = True

    else:
        # CdS or CdSe - polynomial evaluation
        l = peak_wl
        steps = [
            f"Step 1: Identify peak wavelength",
            f"  λ_peak = {l:.2f} nm",
            "",
            f"Step 2: Apply sizing formula (polynomial evaluation)",
            f"  {formula_info['formula_str']}",
        ]

        if qd_type.lower() == "cds":
            steps += [
                f"  D = –6.6521×10⁻⁸ × ({l:.1f})³ + 1.9557×10⁻⁴ × ({l:.1f})²",
                f"      – 9.2352×10⁻² × ({l:.1f}) + 13.29",
                "",
                f"Step 3: Evaluate each term",
                f"  Term 1 = –6.6521×10⁻⁸ × {l**3:.2f} = {-6.6521e-8 * l**3:.4f}",
                f"  Term 2 = 1.9557×10⁻⁴ × {l**2:.2f} = {1.9557e-4 * l**2:.4f}",
                f"  Term 3 = –9.2352×10⁻² × {l:.1f} = {-9.2352e-2 * l:.4f}",
                f"  Term 4 = 13.29",
            ]
        else:
            steps += [
                f"  D = 1.6122×10⁻⁹ × ({l:.1f})⁴ – 2.6575×10⁻⁶ × ({l:.1f})³",
                f"      + 1.6242×10⁻³ × ({l:.1f})² – 0.4277 × ({l:.1f}) + 41.57",
                "",
                f"Step 3: Evaluate each term",
                f"  Term 1 = 1.6122×10⁻⁹ × {l**4:.2f} = {1.6122e-9 * l**4:.4f}",
                f"  Term 2 = –2.6575×10⁻⁶ × {l**3:.2f} = {-2.6575e-6 * l**3:.4f}",
                f"  Term 3 = 1.6242×10⁻³ × {l**2:.2f} = {1.6242e-3 * l**2:.4f}",
                f"  Term 4 = –0.4277 × {l:.1f} = {-0.4277 * l:.4f}",
                f"  Term 5 = 41.57",
            ]

        steps += [
            "",
            f"Step 4: Sum all terms",
            f"  D = {size_nm:.4f} nm",
        ]

        for step in steps:
            p = doc.add_paragraph(step, style="Normal")
            if step.startswith("Step"):
                p.runs[0].bold = True

    doc.add_paragraph()

    # --- Size Result ---
    doc.add_heading("3.3 Final Result", level=2)

    result_table = doc.add_table(rows=4, cols=2, style="Light Grid Accent 1")
    result_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    result_data = [
        (f"{qd_name} QD Diameter", f"{size_nm:.3f} nm"),
        ("Peak Wavelength", f"{peak_wl:.2f} nm"),
        ("Peak Energy", f"{1240/peak_wl:.3f} eV"),
        ("Formula Reference", formula_info["reference"]),
    ]
    for i, (label, value) in enumerate(result_data):
        result_table.rows[i].cells[0].text = label
        result_table.rows[i].cells[1].text = str(value)
        for cell in result_table.rows[i].cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

    # --- Section 4: HWHM Analysis & Size Distribution ---
    doc.add_heading("4. First Exciton Peak Linewidth & Size Distribution", level=1)

    doc.add_paragraph(
        "The half width at half maximum (HWHM) of the first exciton absorption "
        "peak provides information about the size distribution (polydispersity) "
        "of the quantum dot ensemble. The broadening of the exciton peak arises "
        "from inhomogeneous broadening due to the finite size distribution."
    )

    doc.add_heading("4.1 Gaussian Fitting of the First Exciton Peak", level=2)

    if gaussian_fit and gaussian_fit["fit_success"]:
        params = gaussian_fit["fit_params"]
        doc.add_paragraph(
            f"The first exciton absorption peak was fitted with a Gaussian function:"
        )
        fit_eq = doc.add_paragraph()
        fit_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
        eq_run = fit_eq.add_run(
            f"A(λ) = A₀ · exp(-(λ - λ₀)² / (2σ_g²)) + baseline"
        )
        eq_run.font.italic = True
        eq_run.font.size = Pt(11)

        doc.add_paragraph()
        gf_table = doc.add_table(rows=7, cols=2, style="Light Grid Accent 1")
        gf_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        gf_data = [
            ("Peak Center (λ₀)", f"{params['x0']:.2f} nm"),
            ("Peak Amplitude (A₀)", f"{params['A']:.4f} a.u."),
            ("Gaussian Width (σ_g, wavelength)", f"{params['sigma']:.2f} nm"),
            ("Baseline", f"{params['baseline']:.4f} a.u."),
            ("FWHM = 2√(2ln2)·σ_g", f"{gaussian_fit['fwhm_nm']:.2f} nm ({gaussian_fit['fwhm_eV']*1000:.2f} meV)"),
            ("HWHM = √(2ln2)·σ_g", f"{gaussian_fit['hwhm_nm']:.2f} nm ({gaussian_fit['hwhm_eV']*1000:.2f} meV)"),
            ("Fit R²", f"{gaussian_fit['fit_r2']:.4f}"),
        ]
        for i, (label, value) in enumerate(gf_data):
            gf_table.rows[i].cells[0].text = label
            gf_table.rows[i].cells[1].text = str(value)
            for cell in gf_table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(2)
                    paragraph.paragraph_format.space_after = Pt(2)
    else:
        doc.add_paragraph(
            "Gaussian fitting of the first exciton peak was attempted but did not "
            "converge satisfactorily. HWHM could not be reliably determined."
        ).italic = True

    doc.add_paragraph()

    # Section 4.2: Size Distribution
    doc.add_heading("4.2 Size Distribution (σ) from Inhomogeneous Broadening", level=2)

    doc.add_paragraph(
        "The relationship between the HWHM of the first exciton peak and the "
        "size distribution width σ is derived from error propagation on the "
        "empirical sizing curve E(d):"
    )

    # Formula
    sigma_formula = doc.add_paragraph()
    sigma_formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sf_run = sigma_formula.add_run("σ_d = HWHM(E) / |dE/dd|_{d=d_mean}")
    sf_run.font.italic = True
    sf_run.font.size = Pt(12)
    sf_run.font.bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        "This approach assumes that the broadening of the first exciton peak is "
        "dominated by inhomogeneous broadening from the size distribution, and "
        "that the sizing curve E(d) is approximately linear over the width of "
        "the distribution."
    )
    doc.add_paragraph(
        "Reference: Nikolaev & Averkiev, Appl. Phys. Lett. 95, 263107 (2009); "
        "Wu et al., Appl. Phys. Lett. 51, 710 (1987)."
    )

    if sigma_info and not np.isnan(sigma_info.get("sigma_nm", np.nan)):
        doc.add_paragraph()
        doc.add_heading("4.2.1 Derivative of the Sizing Curve", level=3)

        deriv_info = SIZE_FORMULAS.get(qd_type.lower(), {})
        doc.add_paragraph(
            f"The derivative dE/dd is evaluated at the mean diameter d = {size_nm:.4f} nm "
            f"using the {deriv_info.get('name', qd_type)} sizing curve:"
        )

        deriv_str = deriv_info.get("deriv_str", "")
        if deriv_str:
            dp = doc.add_paragraph()
            dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            dr = dp.add_run(deriv_str)
            dr.font.italic = True
            dr.font.size = Pt(10)

        doc.add_paragraph()

        # Step by step calculation
        doc.add_heading("4.2.2 Calculation Steps", level=3)

        if qd_type.lower() in ("pbs", "pbse"):
            steps = [
                f"Step 1: Measured HWHM from Gaussian fit",
                f"  HWHM = {gaussian_fit['hwhm_eV']*1000:.2f} meV = {gaussian_fit['hwhm_eV']:.5f} eV",
                "",
                f"Step 2: Calculate |dE/dd| at d = {size_nm:.4f} nm",
            ]
            if qd_type.lower() == "pbs":
                a, b = 0.0252, 0.283
                u = a * size_nm**2 + b * size_nm
                numerator = 2 * a * size_nm + b
                steps += [
                    f"  u = 0.0252·({size_nm:.4f})² + 0.283·({size_nm:.4f}) = {u:.6f}",
                    f"  dE/dd = -({numerator:.6f}) / ({u:.6f})² = -{numerator/u**2:.6f} eV/nm",
                    f"  |dE/dd| = {abs(numerator/u**2):.6f} eV/nm",
                ]
            else:
                a, b, c_val = 0.016, 0.209, 0.45
                u = a * size_nm**2 + b * size_nm + c_val
                numerator = 2 * a * size_nm + b
                steps += [
                    f"  u = 0.016·({size_nm:.4f})² + 0.209·({size_nm:.4f}) + 0.45 = {u:.6f}",
                    f"  dE/dd = -({numerator:.6f}) / ({u:.6f})² = -{numerator/u**2:.6f} eV/nm",
                    f"  |dE/dd| = {abs(numerator/u**2):.6f} eV/nm",
                ]
        else:
            l = peak_info["peak_wavelength"]
            steps = [
                f"Step 1: Measured HWHM from Gaussian fit",
                f"  HWHM = {gaussian_fit['hwhm_eV']*1000:.2f} meV = {gaussian_fit['hwhm_eV']:.5f} eV",
                "",
                f"Step 2: Calculate dD/dλ at λ = {l:.2f} nm",
            ]
            if qd_type.lower() == "cds":
                dD_dlam = -1.99563e-7 * l**2 + 3.9114e-4 * l - 9.2352e-2
                steps += [
                    f"  dD/dλ = -1.9956×10⁻⁷·λ² + 3.9114×10⁻⁴·λ - 9.2352×10⁻²",
                    f"  dD/dλ = {dD_dlam:.6f} nm/nm",
                ]
            else:
                dD_dlam = 6.4488e-9 * l**3 - 7.9725e-6 * l**2 + 3.2484e-3 * l - 0.4277
                steps += [
                    f"  dD/dλ = 6.4488×10⁻⁹·λ³ - 7.9725×10⁻⁶·λ² + 3.2484×10⁻³·λ - 0.4277",
                    f"  dD/dλ = {dD_dlam:.6f} nm/nm",
                ]
            steps += [
                "",
                f"Step 3: Calculate dE/dD via chain rule",
                f"  dE/dλ = -1240/λ² = -1240/({l:.2f})² = {-1240/l**2:.6f} eV/nm",
                f"  dE/dD = (dE/dλ) / (dD/dλ) = {-1240/l**2:.6f} / {dD_dlam:.6f}",
                f"  dE/dD = {sigma_info['deriv_dE_dd']:.6f} eV/nm",
            ]

        steps += [
            "",
            f"Step 4: Calculate size distribution σ_d",
            f"  σ_d = HWHM / |dE/dd|",
            f"  σ_d = {gaussian_fit['hwhm_eV']:.5f} / {sigma_info['deriv_dE_dd']:.6f}",
            f"  σ_d = {sigma_info['sigma_nm']:.4f} nm",
            "",
            f"Step 5: Relative size distribution",
            f"  σ_d / d = {sigma_info['sigma_nm']:.4f} / {size_nm:.4f}",
            f"  Relative σ = {sigma_info['relative_sigma_percent']:.2f}%",
        ]

        for step in steps:
            p = doc.add_paragraph(step, style="Normal")
            if step.startswith("Step"):
                p.runs[0].bold = True

        doc.add_paragraph()

        # Result table
        doc.add_heading("4.2.3 Size Distribution Results", level=3)
        sd_table = doc.add_table(rows=6, cols=2, style="Light Grid Accent 1")
        sd_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        sd_data = [
            ("Mean Diameter (d)", f"{size_nm:.4f} nm"),
            ("HWHM (energy)", f"{gaussian_fit['hwhm_eV']*1000:.2f} meV"),
            ("|dE/dd| at d_mean", f"{sigma_info['deriv_dE_dd']:.6f} eV/nm"),
            ("Size Distribution (σ_d)", f"{sigma_info['sigma_nm']:.4f} nm"),
            ("Relative σ (σ_d / d)", f"{sigma_info['relative_sigma_percent']:.2f}%"),
            ("FWHM Size Range (±2σ covers 95%)", f"{size_nm - 2*sigma_info['sigma_nm']:.2f} – {size_nm + 2*sigma_info['sigma_nm']:.2f} nm"),
        ]
        for i, (label, value) in enumerate(sd_data):
            sd_table.rows[i].cells[0].text = label
            sd_table.rows[i].cells[1].text = str(value)
            for cell in sd_table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(2)
                    paragraph.paragraph_format.space_after = Pt(2)
    else:
        doc.add_paragraph(
            "Size distribution could not be calculated because the Gaussian "
            "fitting did not converge or the sizing curve derivative could not "
            "be evaluated."
        ).italic = True

    doc.add_paragraph()

    # --- Section 5: UV-Vis Spectrum ---
    doc.add_heading("5. UV-Vis Absorption Spectrum", level=1)

    # 5.1 Baseline subtraction equation
    doc.add_heading("5.1 Baseline Subtraction", level=2)
    if gaussian_fit and gaussian_fit["fit_success"]:
        baseline_model = gaussian_fit.get("baseline_model", "constant")
        bl_params = gaussian_fit["fit_params"].get("baseline_params", {})
        bl_at_peak = gaussian_fit["fit_params"].get("baseline", 0)

        if baseline_model == "constant":
            b = bl_params.get("b", 0)
            doc.add_paragraph(
                f"The baseline was modeled as a constant offset:\n"
                f"    Baseline(λ) = {b:.4f}\n"
                f"Baseline at first exciton peak ({peak_info['peak_wavelength']:.1f} nm): {bl_at_peak:.4f} a.u."
            )
        elif baseline_model == "linear":
            a = bl_params.get("a", 0)
            b = bl_params.get("b", 0)
            doc.add_paragraph(
                f"The baseline was modeled as a linear function:\n"
                f"    Baseline(λ) = a·λ + b = {a:.6f}·λ + {b:.4f}\n"
                f"Baseline at first exciton peak ({peak_info['peak_wavelength']:.1f} nm): {bl_at_peak:.4f} a.u."
            )
        else:  # exponential
            c = bl_params.get("c", 0)
            d = bl_params.get("d", 0)
            b0 = bl_params.get("b0", 0)
            doc.add_paragraph(
                f"The baseline was modeled as an exponential decay (suitable for absorption edges):\n"
                f"    Baseline(λ) = c·exp(-d·(λ - λ_ref)) + b₀\n"
                f"    c = {c:.4f}, d = {d:.6f} nm⁻¹, b₀ = {b0:.2e}\n"
                f"Baseline at first exciton peak ({peak_info['peak_wavelength']:.1f} nm): {bl_at_peak:.4f} a.u.\n"
                f"The baseline was fitted to spectral regions outside the exciton peak core "
                f"(two-step method: baseline from peak-excluded data, then pure Gaussian on subtracted data)."
            )
    else:
        doc.add_paragraph("No baseline information available (Gaussian fit did not converge).")

    doc.add_paragraph()

    # 5.2 Plot description
    doc.add_heading("5.2 Figure Legend", level=2)
    doc.add_paragraph(
        "The figure below displays the following traces:"
    )

    # Build legend items
    legend_items = [
        ("Blue solid line", "Absorption", "Raw UV-Vis absorption spectrum as measured by the spectrophotometer."),
        ("Vermilion/red solid line", "Pure Gauss Fit",
         "Pure Gaussian function fitted to the baseline-subtracted data over the first exciton peak region. "
         "This curve represents the exciton absorption after removing the background (baseline)."),
        ("Green dashed line", "HWHM",
         "Half Width at Half Maximum (HWHM) of the fitted Gaussian. "
         "The horizontal bar is drawn at half the Gaussian amplitude (A/2). "
         "The separation between the two vertical end-ticks equals 2 × HWHM = FWHM."),
        ("Red dot", "First Exciton Peak",
         f"Position of the first exciton absorption peak at {peak_info['peak_wavelength']:.2f} nm "
         f"({peak_info.get('peak_energy_eV', peak_info.get('peak_energy', 0)):.4f} eV)."),
    ]

    for label, name, desc in legend_items:
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        p.add_run(f"{name} — {desc}")

    doc.add_paragraph()
    doc.add_paragraph(
        "The baseline (absorption edge or scattering background) is mathematically subtracted "
        "before Gaussian fitting. The pure Gaussian curve represents the homogeneous and "
        "inhomogeneous broadening of the first exciton transition, from which the HWHM "
        "and size distribution are derived."
    )

    doc.add_paragraph()

    if os.path.exists(plot_path):
        doc.add_picture(plot_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # --- Section 6: Method Summary ---
    doc.add_heading("6. Method Summary", level=1)
    doc.add_paragraph(
        "This analysis determines quantum dot sizes from UV-Vis absorption spectra "
        "using the following procedure:"
    )
    methods = [
        "1. Raw UV-Vis data is parsed from the instrument output file.",
        "2. The spectrum is smoothed using a Savitzky-Golay filter to reduce noise.",
        "3. Absorption peaks are detected and the first exciton peak is identified "
        "as the most prominent peak at the longest wavelength (lowest energy).",
        f"4. The peak wavelength is converted to energy: E₀ = 1240 / λ.",
        f"5. The empirical sizing formula from {formula_info['reference'].split(',')[0]} "
        f"is applied to calculate the quantum dot diameter d.",
        "6. A Gaussian function is fitted to the first exciton peak region to extract "
        "the half width at half maximum (HWHM).",
        "7. The size distribution width σ_d is calculated via error propagation: "
        "σ_d = HWHM(E) / |dE/dd|, where dE/dd is the derivative of the sizing curve "
        "at the mean diameter.",
        "8. Results are visualized in a publication-quality plot (with Gaussian fit "
        "overlay and HWHM markers) and compiled into this analysis report.",
    ]
    for m in methods:
        doc.add_paragraph(m, style="Normal")

    doc.add_paragraph()

    # --- Warnings (if any) ---
    if warnings:
        doc.add_heading("⚠ Validity Warnings", level=2)
        for w in warnings:
            p = doc.add_paragraph(w, style="Normal")
            p.runs[0].font.color.rgb = RGBColor(180, 30, 30) if p.runs else None
        doc.add_paragraph(
            "The empirical sizing formula was calibrated using TEM measurements "
            "over a specific wavelength/size range. Results outside this range "
            "should be treated with caution and ideally verified by TEM or XRD."
        ).italic = True
        doc.add_paragraph()

    # --- Disclaimer ---
    doc.add_heading("Disclaimer", level=2)
    doc.add_paragraph(
        "The sizing formulas used are empirical calibrations based on TEM measurements "
        "and tight-binding calculations. Accuracy depends on: (1) correct identification "
        "of the first exciton peak, (2) the quantum dots being within the calibrated "
        "size range, (3) the synthetic method being comparable to those used in the "
        "reference studies. For critical applications, verify sizes with TEM or XRD."
    ).italic = True

    # Save
    safe_name = re.sub(r'[^\w\s-]', '', data["sample_name"]).strip()
    safe_name = re.sub(r'[-\s]+', '_', safe_name)
    doc_path = os.path.join(output_dir, f"UV_Analysis_Report_{safe_name}.docx")
    doc.save(doc_path)

    return doc_path


# ============================================================
# MAIN ENTRY POINT
# ============================================================

def analyze_single_file(filepath, qd_type, output_dir=None, smooth_window=15,
                        x_range=None, y_range=None, output_formats=None,
                        dpi=300, figure_width="double", show_annotation=False,
                        fit_range_nm=None, baseline_mode="auto"):
    """Analyze a single UV-Vis data file."""
    if output_dir is None:
        output_dir = os.path.dirname(filepath) or "."

    os.makedirs(output_dir, exist_ok=True)

    if output_formats is None:
        output_formats = ["png"]

    # Parse data
    print(f"Parsing: {filepath}")
    data = parse_uv_txt(filepath)
    print(f"  Sample: {data['sample_name']}")
    print(f"  Data points: {len(data['wavelength'])}")
    print(f"  Wavelength range: {data['wavelength'].min():.1f} – {data['wavelength'].max():.1f} nm")

    # Find peak
    print(f"  Finding first exciton peak for {qd_type.upper()}...")
    peak_info = find_first_exciton_peak(
        data["wavelength"], data["absorbance"], qd_type,
        smooth_window=smooth_window
    )
    print(f"  Peak: {peak_info['peak_wavelength']:.2f} nm "
          f"({peak_info['peak_energy_eV']:.3f} eV)")

    # Calculate size
    print(f"  Calculating size...")
    warnings_list = []
    try:
        size_nm, E0, warnings_list = calculate_size(qd_type, peak_info["peak_wavelength"])
        print(f"  Size: {size_nm:.3f} nm")
        for w in warnings_list:
            print(f"  {w}")
    except ValueError as e:
        print(f"  ERROR: {e}")
        size_nm = float("nan")

    # Gaussian fitting for HWHM
    print(f"  Fitting Gaussian to first exciton peak...")
    peak_idx = peak_info.get("first_exciton_details", {}).get("index",
                    int(np.argmin(np.abs(data["wavelength"] - peak_info["peak_wavelength"]))))
    gfit = fit_gaussian_to_peak(data["wavelength"], data["absorbance"], peak_idx, qd_type=qd_type, fit_range_nm=fit_range_nm, baseline_mode=baseline_mode)
    if gfit["fit_success"]:
        print(f"  HWHM: {gfit['hwhm_nm']:.2f} nm ({gfit['hwhm_eV']*1000:.2f} meV)")
        print(f"  Gaussian R^2: {gfit['fit_r2']:.4f}")
        if gfit.get("hwhm_warning"):
            print(f"  WARNING: {gfit['hwhm_warning']}")
            warnings_list.append(gfit["hwhm_warning"])
    else:
        print(f"  WARNING: Gaussian fitting failed, HWHM not available")

    # Calculate size distribution sigma
    sigma_info = None
    if gfit["fit_success"] and not np.isnan(size_nm):
        print(f"  Calculating size distribution...")
        sigma_info = calculate_size_distribution(
            qd_type, size_nm, gfit["hwhm_eV"], peak_info["peak_wavelength"]
        )
        if not np.isnan(sigma_info["sigma_nm"]):
            print(f"  sigma_d = {sigma_info['sigma_nm']:.4f} nm "
                  f"({sigma_info['relative_sigma_percent']:.1f}%)")
        else:
            print(f"  WARNING: Could not calculate sigma")

    # Create plot
    print(f"  Generating UV-Vis plot...")
    plot_paths = create_uv_plot(data, peak_info, qd_type, size_nm, output_dir,
                                gaussian_fit=gfit, sigma_info=sigma_info,
                                x_range=x_range, y_range=y_range,
                                output_formats=output_formats, dpi=dpi,
                                figure_width=figure_width,
                                show_annotation=show_annotation)
    png_path = plot_paths.get("png", next(iter(plot_paths.values())))
    for fmt, p in plot_paths.items():
        print(f"  Plot ({fmt}): {p}")

    # Create Word report (embed PNG)
    print(f"  Generating Word report...")
    doc_path = create_word_report(data, peak_info, qd_type, size_nm, png_path, output_dir,
                                   warnings=warnings_list, gaussian_fit=gfit,
                                   sigma_info=sigma_info)
    if doc_path:
        print(f"  Report saved: {doc_path}")

    return {
        "data": data,
        "peak_info": peak_info,
        "size_nm": size_nm,
        "gaussian_fit": gfit,
        "sigma_info": sigma_info,
        "plot_paths": plot_paths,
        "plot_path": png_path,
        "doc_path": doc_path,
        "warnings": warnings_list,
    }


def analyze_multiple_files(filepaths, qd_type, output_dir=None):
    """Analyze multiple UV-Vis data files."""
    results = []
    for fp in filepaths:
        try:
            if output_dir is None:
                out_dir = os.path.dirname(fp) or "."
            else:
                out_dir = output_dir
            result = analyze_single_file(fp, qd_type, out_dir)
            results.append(result)
        except Exception as e:
            print(f"ERROR processing {fp}: {e}")
            results.append({"error": str(e), "filepath": fp})
    return results


def detect_qd_type_from_filename(filename):
    """Try to detect QD type from filename."""
    fname = filename.lower()
    if "pbs" in fname:
        return "pbs"
    elif "pbse" in fname:
        return "pbse"
    elif "cdse" in fname:
        return "cdse"
    elif "cds" in fname:
        return "cds"
    return None


def infer_qd_type_from_spectrum(filepath, smooth_window=15):
    """
    Parse the data, find peaks, and infer the most likely QD type
    based on where the first exciton peak falls relative to each
    material's calibrated wavelength range.

    Returns a ranked list of (qd_type, confidence, reason) tuples.
    Confidence: 'high' = peak inside range, 'medium' = peak near range
                edge, 'low' = peak outside range but closest match.
    """
    data = parse_uv_txt(filepath)
    wl = data["wavelength"]
    abs_data = data["absorbance"]

    # Find the most prominent peak
    # Use a simple approach: find the global maximum in the smoothed spectrum
    n_points = len(abs_data)
    actual_window = min(smooth_window if smooth_window % 2 == 1 else smooth_window + 1,
                        n_points - 2 if (n_points - 2) % 2 == 1 else n_points - 3)
    if actual_window >= 5:
        smoothed = savgol_filter(abs_data, actual_window, 3)
    else:
        smoothed = abs_data

    abs_range = np.max(smoothed) - np.min(smoothed)
    height_threshold = max(np.min(smoothed) + 0.05 * abs_range, 0.0)
    peaks, properties = find_peaks(
        smoothed, height=height_threshold, prominence=0.03 * abs_range, distance=5
    )

    if len(peaks) == 0:
        peak_idx = np.argmax(smoothed)
    else:
        # Pick longest-wavelength significant peak
        max_prom = max(properties.get("prominences", [1]))
        sig_peaks = [(i, p) for i, p in enumerate(peaks)
                     if properties["prominences"][i] >= 0.1 * max_prom]
        if sig_peaks:
            # sig_peaks contains (index_in_peaks, wavelength_index)
            # Pick the one with the longest wavelength (largest wl value)
            best = max(sig_peaks, key=lambda x: wl[x[1]])
            peak_idx = best[1]  # wavelength index (direct index into wl array)
        else:
            peak_idx = peaks[np.argmax([wl[p] for p in peaks])]

    peak_wl = wl[peak_idx]
    peak_eV = 1240.0 / peak_wl if peak_wl > 0 else 0

    # Score each QD type
    results = []
    for qd_type, r in VALID_RANGES.items():
        in_range = r["wl_min"] <= peak_wl <= r["wl_max"]
        range_width = r["wl_max"] - r["wl_min"]
        range_center = (r["wl_min"] + r["wl_max"]) / 2

        # Normalized distance: how far (in range-widths) from nearest edge
        if in_range:
            margin_norm = 0.0
        elif peak_wl < r["wl_min"]:
            margin_norm = (r["wl_min"] - peak_wl) / range_width
        else:
            margin_norm = (peak_wl - r["wl_max"]) / range_width

        name = SIZE_FORMULAS[qd_type]["name"]

        if in_range:
            # Further distinction: how centered within the range
            centering = abs(peak_wl - range_center) / (range_width / 2)
            if centering < 0.3:
                confidence = "high"
            elif centering < 0.7:
                confidence = "high"
            else:
                confidence = "medium"
            reason = (f"Peak at {peak_wl:.0f} nm ({peak_eV:.3f} eV) falls within "
                      f"the {name} calibrated range ({r['wl_min']:.0f}-{r['wl_max']:.0f} nm)")
        elif margin_norm < 0.15:
            confidence = "medium"
            direction = "below" if peak_wl < r["wl_min"] else "above"
            reason = (f"Peak at {peak_wl:.0f} nm is near the {name} calibrated range "
                      f"({r['wl_min']:.0f}-{r['wl_max']:.0f} nm, {direction} by "
                      f"{abs(peak_wl - r['wl_min'] if peak_wl < r['wl_min'] else peak_wl - r['wl_max']):.0f} nm)")
        elif margin_norm < 0.5:
            confidence = "low"
            direction = "below" if peak_wl < r["wl_min"] else "above"
            reason = (f"Peak at {peak_wl:.0f} nm is {direction} the "
                      f"{name} calibrated range ({r['wl_min']:.0f}-{r['wl_max']:.0f} nm)")
        else:
            confidence = "low"
            reason = (f"Peak at {peak_wl:.0f} nm is far outside the "
                      f"{name} calibrated range ({r['wl_min']:.0f}-{r['wl_max']:.0f} nm)")

        results.append({
            "qd_type": qd_type,
            "name": name,
            "confidence": confidence,
            "reason": reason,
            "peak_wl": peak_wl,
            "peak_eV": peak_eV,
            "margin_norm": margin_norm,
        })

    # Sort by confidence (high > medium > low), then by normalized margin
    conf_order = {"high": 0, "medium": 1, "low": 2}
    results.sort(key=lambda x: (conf_order[x["confidence"]], x["margin_norm"]))

    return results, data


def analyze_to_dict(filepath, qd_type, output_dir=None, smooth_window=15,
                    x_range=None, y_range=None, output_formats=None,
                    dpi=300, figure_width="double", show_annotation=False,
                    fit_range_nm=None, baseline_mode="auto"):
    """
    Run full analysis and return a JSON-serializable dict for web API.
    Wraps analyze_single_file(), converts numpy arrays to lists, NaN to null.
    """
    result = analyze_single_file(filepath, qd_type, output_dir, smooth_window,
                                 x_range=x_range, y_range=y_range,
                                 output_formats=output_formats, dpi=dpi,
                                 figure_width=figure_width,
                                 show_annotation=show_annotation,
                                 fit_range_nm=fit_range_nm,
                                 baseline_mode=baseline_mode)

    data = result["data"]
    peak_info = result["peak_info"]
    gfit = result["gaussian_fit"]
    sigma_info = result["sigma_info"]

    # Build fit curve arrays for Plotly overlay
    gaussian_fit_wl = []
    gaussian_fit_abs = []
    gauss_only_wl = []
    gauss_only_abs = []
    subtracted_wl = []
    subtracted_abs = []
    if gfit["fit_success"] and gfit.get("fit_wavelength") is not None:
        gaussian_fit_wl = gfit["fit_wavelength"].tolist()
        gaussian_fit_abs = gfit["fit_absorbance"].tolist()
        if gfit.get("gauss_only") is not None:
            gauss_only_wl = gfit["fit_wavelength"].tolist()
            gauss_only_abs = gfit["gauss_only"].tolist()
        if gfit.get("fit_wl_data") is not None and gfit.get("abs_subtracted") is not None:
            subtracted_wl = gfit["fit_wl_data"].tolist()
            subtracted_abs = gfit["abs_subtracted"].tolist()

    # HWHM markers for Plotly
    peak_wl = float(peak_info["peak_wavelength"])
    hwhm_left_wl = None
    hwhm_right_wl = None
    hwhm_half_max = None
    if gfit["fit_success"] and not np.isnan(gfit.get("hwhm_nm", float("nan"))):
        hwhm_left_wl = peak_wl - float(gfit["hwhm_nm"])
        hwhm_right_wl = peak_wl + float(gfit["hwhm_nm"])
        # Half-max absorbance: baseline_at_peak + A/2
        if gfit["fit_params"]["A"] is not None:
            hwhm_half_max = float(gfit["fit_params"]["A"]) / 2.0
        else:
            hwhm_half_max = float(peak_info["peak_absorbance"]) / 2.0

    # Get formula info
    formula_info = SIZE_FORMULAS.get(qd_type, {})
    valid_range = VALID_RANGES.get(qd_type, {})

    def safe_float(v):
        """Convert to float, return None for NaN."""
        if v is None:
            return None
        f = float(v)
        return None if np.isnan(f) else f

    # Serialize warnings
    from numpy import ndarray
    def json_safe(obj):
        """Recursively convert numpy types to Python native types."""
        if isinstance(obj, (np.integer,)):
            return int(obj)
        elif isinstance(obj, (np.floating,)):
            val = float(obj)
            return None if np.isnan(val) else val
        elif isinstance(obj, ndarray):
            return obj.tolist()
        elif isinstance(obj, dict):
            return {k: json_safe(v) for k, v in obj.items()}
        elif isinstance(obj, list):
            return [json_safe(i) for i in obj]
        return obj

    response = {
        "sample_name": str(data["sample_name"]),
        "filename": str(data["filename"]),
        "qd_type": qd_type,
        "peak_wavelength_nm": safe_float(peak_info["peak_wavelength"]),
        "peak_energy_eV": safe_float(peak_info["peak_energy_eV"]),
        "peak_absorbance": safe_float(peak_info["peak_absorbance"]),
        "size_nm": safe_float(result["size_nm"]),
        "hwhm_nm": safe_float(gfit.get("hwhm_nm")),
        "hwhm_eV": safe_float(gfit.get("hwhm_eV")),
        "fwhm_nm": safe_float(gfit.get("fwhm_nm")),
        "fwhm_eV": safe_float(gfit.get("fwhm_eV")),
        "gaussian_sigma_nm": safe_float(gfit.get("gaussian_sigma_nm")),
        "fit_r2": safe_float(gfit.get("fit_r2")),
        "fit_success": bool(gfit.get("fit_success", False)),
        "hwhm_warning": gfit.get("hwhm_warning"),
        "hwhm_min_meV": gfit.get("hwhm_min_meV"),
        "hwhm_max_meV": gfit.get("hwhm_max_meV"),
        "sigma_nm": safe_float(sigma_info.get("sigma_nm")) if sigma_info else None,
        "relative_sigma_percent": safe_float(sigma_info.get("relative_sigma_percent")) if sigma_info else None,
        "deriv_dE_dd": safe_float(sigma_info.get("deriv_dE_dd")) if sigma_info else None,
        "formula_name": formula_info.get("name", ""),
        "formula_str": formula_info.get("formula_str", ""),
        "deriv_str": formula_info.get("deriv_str", ""),
        "reference": formula_info.get("reference", ""),
        "valid_range_wl": f"{valid_range.get('wl_min', '?')}-{valid_range.get('wl_max', '?')} nm",
        "valid_range_size": f"{valid_range.get('size_min', '?')}-{valid_range.get('size_max', '?')} nm",
        "wavelength": data["wavelength"].tolist(),
        "absorbance": data["absorbance"].tolist(),
        "gaussian_fit_wl": gaussian_fit_wl,
        "gaussian_fit_abs": gaussian_fit_abs,
        "gauss_only_wl": gauss_only_wl,
        "gauss_only_abs": gauss_only_abs,
        "subtracted_wl": subtracted_wl,
        "subtracted_abs": subtracted_abs,
        "baseline_model": gfit.get("baseline_model"),
        "hwhm_left_wl": hwhm_left_wl,
        "hwhm_right_wl": hwhm_right_wl,
        "hwhm_half_max_abs": hwhm_half_max,
        "plot_path": result.get("plot_path", ""),
        "plot_paths": {fmt: str(p) for fmt, p in result.get("plot_paths", {}).items()},
        "doc_path": result.get("doc_path", ""),
        "warnings": [str(w) for w in (result.get("warnings") or [])],
        "all_peaks": [{
            "wavelength_nm": safe_float(p.get("wavelength")),
            "absorbance": safe_float(p.get("absorbance")),
            "energy_eV": safe_float(p.get("energy_eV")),
        } for p in peak_info.get("all_peaks", [])],
    }

    return response


def get_info():
    """Return formula metadata for the web frontend."""
    import copy
    formulas_out = {}
    for qd_type, info in SIZE_FORMULAS.items():
        formulas_out[qd_type] = {
            "name": info["name"],
            "reference": info["reference"],
            "formula_str": info["formula_str"],
            "deriv_str": info["deriv_str"],
            "variable": info["variable"],
            "method": info["method"],
            "valid_range": info["valid_range"],
            "bulk_Eg_eV": info["bulk_bandgap_eV"],
        }
    return {
        "formulas": formulas_out,
        "valid_ranges": dict(VALID_RANGES),
    }


def main():
    parser = argparse.ArgumentParser(
        description="UV-Vis Quantum Dot Size Analysis",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python uv_analysis.py data.txt --type pbs
  python uv_analysis.py *.txt --type cdse
  python uv_analysis.py data.txt --type pbs --output ./results
  python uv_analysis.py data.txt --type auto  (detect from filename)
  python uv_analysis.py data.txt --type pbs --xmin 800 --xmax 1800 --ymin 0
  python uv_analysis.py data.txt --type pbs --format pdf,svg,png --dpi 600
  python uv_analysis.py data.txt --type pbs --figure-width single
        """
    )
    parser.add_argument("files", nargs="+", help="UV-Vis data file(s) to analyze")
    parser.add_argument("--type", "-t", dest="qd_type", required=True,
                        choices=["pbs", "pbse", "cds", "cdse", "auto"],
                        help="Quantum dot type (or 'auto' to detect from filename)")
    parser.add_argument("--output", "-o", dest="output_dir", default=None,
                        help="Output directory for plots and reports")
    parser.add_argument("--smooth", "-s", dest="smooth_window", type=int, default=15,
                        help="Savitzky-Golay smoothing window (odd number, default: 15)")

    # New: Axis range control
    axis_group = parser.add_argument_group("Axis range control")
    axis_group.add_argument("--xmin", type=float, default=None,
                            help="Custom X-axis minimum (nm)")
    axis_group.add_argument("--xmax", type=float, default=None,
                            help="Custom X-axis maximum (nm)")
    axis_group.add_argument("--ymin", type=float, default=None,
                            help="Custom Y-axis minimum (absorbance)")
    axis_group.add_argument("--ymax", type=float, default=None,
                            help="Custom Y-axis maximum (absorbance)")

    # New: Output format & quality
    output_group = parser.add_argument_group("Output format & quality")
    output_group.add_argument("--format", "-f", dest="output_formats", default="png",
                              help="Output formats: png, pdf, svg, or comma-separated list "
                                   "(e.g. 'png,pdf,svg' or 'all')")
    output_group.add_argument("--dpi", type=int, default=300, choices=[150, 300, 600],
                              help="PNG resolution in DPI (default: 300)")
    output_group.add_argument("--figure-width", dest="figure_width", default="double",
                              choices=["single", "double"],
                              help="Figure width: single (8.5 cm) or double (17 cm)")
    output_group.add_argument("--annotation", dest="show_annotation", action="store_true",
                              default=False,
                              help="Show data annotation box on plot (λ, E, d, HWHM, sigma)")
    output_group.add_argument("--fit-range", dest="fit_range", type=float, default=None,
                              help="Gaussian fitting window half-width in nm "
                                   "(auto-selected per QD type if not specified)")
    output_group.add_argument("--baseline", dest="baseline_mode", default="auto",
                              choices=["auto", "constant", "linear", "exponential"],
                              help="Baseline subtraction method (default: auto — "
                                   "try all and select best via AIC)")

    args = parser.parse_args()

    # Parse axis ranges
    x_range = None
    if args.xmin is not None or args.xmax is not None:
        x_range = (args.xmin, args.xmax)
    y_range = None
    if args.ymin is not None or args.ymax is not None:
        y_range = (args.ymin, args.ymax)

    # Parse output formats
    if args.output_formats.lower() == "all":
        output_formats = ["png", "pdf", "svg"]
    else:
        output_formats = [f.strip().lower() for f in args.output_formats.split(",")]

    # Determine QD type
    if args.qd_type == "auto":
        qd_type = detect_qd_type_from_filename(args.files[0])
        if qd_type is not None:
            print(f"Auto-detected QD type from filename: {qd_type.upper()}")
        else:
            # Fall back to spectral inference
            print("Could not detect QD type from filename.")
            print("Analyzing spectrum to infer QD type...")
            print()
            rankings, preview_data = infer_qd_type_from_spectrum(
                args.files[0], smooth_window=args.smooth_window
            )
            print(f"  Detected peak: {rankings[0]['peak_wl']:.0f} nm "
                  f"({rankings[0]['peak_eV']:.3f} eV)")
            print()
            print("  QD Type Ranking (best match first):")
            print("  " + "-" * 55)
            for i, r in enumerate(rankings):
                marker = {0: ">>>", 1: "   ", 2: "   ", 3: "   "}.get(i, "   ")
                conf_mark = {"high": "***", "medium": " **", "low": "  *"}.get(r["confidence"], "")
                print(f"  {marker} {r['name']:6s} [{r['confidence']:6s}] {conf_mark}")
                print(f"       {r['reason']}")
            print("  " + "-" * 55)
            print()

            best = rankings[0]
            if best["confidence"] == "high":
                qd_type = best["qd_type"]
                print(f"Auto-selected: {best['name']} (high confidence match)")
            elif best["confidence"] == "medium":
                qd_type = best["qd_type"]
                print(f"Auto-selected: {best['name']} (moderate confidence — verify manually)")
            else:
                print("No confident match found. The detected peak may not be a")
                print("first exciton transition, or the material may not be one of")
                print("the four supported types (PbS, PbSe, CdS, CdSe).")
                print()
                print("Please specify the QD type manually with --type")
                sys.exit(1)
    else:
        qd_type = args.qd_type

    # Validate files exist
    valid_files = []
    for pattern in args.files:
        import glob
        matched = glob.glob(pattern)
        if not matched:
            print(f"WARNING: No files matching '{pattern}'")
        valid_files.extend(matched)

    if not valid_files:
        print("ERROR: No valid input files found.")
        sys.exit(1)

    print()
    print(f"Analyzing {len(valid_files)} file(s) as {qd_type.upper()} QDs")
    if x_range:
        print(f"  X range: {x_range[0] if x_range[0] is not None else 'auto'} – "
              f"{x_range[1] if x_range[1] is not None else 'auto'} nm")
    if y_range:
        print(f"  Y range: {y_range[0] if y_range[0] is not None else 'auto'} – "
              f"{y_range[1] if y_range[1] is not None else 'auto'}")
    print(f"  Formats: {', '.join(output_formats).upper()} @ {args.dpi} DPI")
    print(f"  Figure width: {args.figure_width}")
    print("=" * 60)

    if len(valid_files) == 1:
        result = analyze_single_file(
            valid_files[0], qd_type,
            output_dir=args.output_dir,
            smooth_window=args.smooth_window,
            x_range=x_range, y_range=y_range,
            output_formats=output_formats, dpi=args.dpi,
            figure_width=args.figure_width,
            show_annotation=args.show_annotation,
            fit_range_nm=args.fit_range,
            baseline_mode=args.baseline_mode,
        )
    else:
        result = analyze_multiple_files(
            valid_files, qd_type,
            output_dir=args.output_dir
        )

    print("=" * 60)
    print("Analysis complete!")

    # Summary
    if isinstance(result, list):
        for r in result:
            if "error" in r:
                print(f"  {r['filepath']}: ERROR - {r['error']}")
            else:
                print(f"  {r['data']['sample_name']}: "
                      f"Peak={r['peak_info']['peak_wavelength']:.1f} nm, "
                      f"Size={r['size_nm']:.3f} nm")
    else:
        print(f"  {result['data']['sample_name']}: "
              f"Peak={result['peak_info']['peak_wavelength']:.1f} nm, "
              f"Size={result['size_nm']:.3f} nm")


if __name__ == "__main__":
    main()
